import base64
import streamlit as st
from docxtpl import DocxTemplate
from docx import Document
import io
import json
import uuid
import pandas as pd
from modules.atas import ia_utils
from modules.atas import authentique_utils
from modules.atas import email_utils
from modules.atas import history_utils
from modules.atas import config
from modules.ui.sidebar import render_sidebar

try:
    import mammoth
    HAS_MAMMOTH = True
except:
    HAS_MAMMOTH = False

# Tenta importar conversor de PDF (Funciona melhor em Windows/Mac locais)
try:
    from docx2pdf import convert
    HAS_PDF_CONVERTER = True
except:
    HAS_PDF_CONVERTER = False

# Configuração da Página
st.set_page_config(page_title="AGIL | ATAs", layout="wide", page_icon="📄")

render_sidebar(active_page="atas")

# CSS para corrigir alinhamentos
st.markdown("""
<style>
    .stButton button[kind="secondary"] { color: red; border-color: red; }
    .block-container { padding-top: 2rem; }
    /* Ajuste para o botão de lixeira ficar alinhado com input de texto */
    div[data-testid="column"] { align-items: end; }
</style>
""", unsafe_allow_html=True)

st.title("📄 AGIL | ATAs")


def render_pdf_preview(pdf_bytes, height=720):
    encoded_pdf = base64.b64encode(pdf_bytes).decode("utf-8")
    st.markdown(
        f'<iframe src="data:application/pdf;base64,{encoded_pdf}" width="100%" height="{height}" type="application/pdf"></iframe>',
        unsafe_allow_html=True,
    )


def extract_docx_preview_blocks(docx_bytes):
    document = Document(io.BytesIO(docx_bytes))
    blocks = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(("paragraph", text))

    for table in document.tables:
        table_rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                table_rows.append(values)
        if table_rows:
            blocks.append(("table", table_rows))

    return blocks


def render_docx_preview(docx_bytes):
    if HAS_MAMMOTH:
        try:
            result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
            html = result.value.strip()
            if html:
                st.markdown(
                    f"""
                    <div style="background: white; color: #111827; padding: 2rem; border-radius: 0.75rem; border: 1px solid #d1d5db; max-height: 720px; overflow: auto;">
                        {html}
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                if result.messages:
                    st.caption("A visualização foi gerada com pequenas limitações de formatação do DOCX.")
                return
        except Exception:
            pass

    blocks = extract_docx_preview_blocks(docx_bytes)
    if not blocks:
        st.warning("Não foi possível extrair conteúdo legível do documento para visualização.")
        return

    st.caption("Visualização simplificada do conteúdo do .docx.")
    for block_type, content in blocks:
        if block_type == "paragraph":
            st.write(content)
        elif block_type == "table":
            st.table(content)


def merge_corrected_pautas(current_pautas, corrected_pautas):
    corrected_by_id = {
        pauta.get("id"): pauta.get("texto", "")
        for pauta in corrected_pautas
        if pauta.get("id")
    }
    merged = []
    for pauta in current_pautas:
        updated = dict(pauta)
        if pauta.get("id") in corrected_by_id:
            updated["texto"] = corrected_by_id[pauta["id"]]
        merged.append(updated)
    return merged


def set_final_document_state(file_bytes, filename, file_type, docx_bytes=None, pdf_bytes=None):
    deadline_iso = authentique_utils.calculate_deadline()
    deadline_display = deadline_iso[:10].split("-")
    deadline_display = f"{deadline_display[2]}/{deadline_display[1]}/{deadline_display[0]}"

    st.session_state['final_file_blob'] = file_bytes
    st.session_state['final_filename'] = filename
    st.session_state['file_type'] = file_type
    st.session_state['final_docx_blob'] = docx_bytes
    st.session_state['final_pdf_blob'] = pdf_bytes
    st.session_state['final_deadline'] = deadline_display

# --- Inicialização de Estado ---
if 'data_store' not in st.session_state:
    st.session_state.data_store = {
        "meta": {},
        "transparencias_data": {}, 
        "attendance_raw": "",
        "final_file_blob": None,
        "final_filename": "ATA_Reuniao.docx", # Default
        "final_deadline": "",
        "final_pdf_blob": None,
        "final_docx_blob": None,
    }

if 'pautas_dinamicas' not in st.session_state:
    st.session_state.pautas_dinamicas = []

# Funções Auxiliares
@st.cache_data
def load_members():
    try:
        with open(config.EMAIL_DB_PATH, "r", encoding="utf-8") as f:
            return sorted(list(json.load(f).keys()))
    except: return []

@st.cache_data
def load_history_context():
    txt, count = history_utils.load_reference_style()
    return txt, count

all_members = load_members()
history_txt, history_count = load_history_context()

def generate_filename():
    """Gera nome padrão: ATA RG - DD/MM.docx"""
    meta = st.session_state.data_store.get("meta", {})
    if "data_obj" in meta:
        date_str = meta["data_obj"].strftime("%d/%m")
        return f"ATA RG - {date_str}.docx" # Retorna docx, converte depois se der
    return "ATA_RG_Generica.docx"

# Tabs
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
    "1. Dados", "2. Transparências", "3. Pautas", 
    "4. Seções Finais", "5. Revisão e Conclusão", 
    "6. Assinatura", "7. Notificação"
])

# --- TAB 1: DADOS ---
with tab1:
    st.header("Dados da Reunião")
    with st.form("form_dados"):
        c1, c2 = st.columns(2)
        gestao = c1.text_input("Gestão", placeholder="Ex: 2025.2")
        data_reuniao = c2.date_input("Data")
        c3, c4 = st.columns(2)
        inicio = c3.time_input("Início")
        fim = c4.time_input("Fim")
        ausentes = st.multiselect("Membros Ausentes:", options=all_members)
        
        if st.form_submit_button("💾 Salvar Metadados"):
            presentes_list = [m for m in all_members if m not in ausentes]
            st.session_state.data_store["meta"] = {
                "gestão": gestao, "data_atual": data_reuniao.strftime("%d/%m/%Y"),
                "data_obj": data_reuniao, # Salva objeto data para formatar nome depois
                "inicio": inicio.strftime("%H:%M"), "fim": fim.strftime("%H:%M"),
                "presentes": ", ".join(presentes_list) + "."
            }
            st.session_state.data_store["attendance_raw"] = "\n".join(presentes_list)
            st.success(f"Dados salvos! {len(presentes_list)} presentes.")

# --- TAB 2: TRANSPARÊNCIAS ---
with tab2:
    st.header("Transparências")
    houve_transp = st.checkbox("Houve Transparências?", value=False)
    if houve_transp:
        c_file, c_act = st.columns([4, 1], vertical_alignment="center") # Alinhamento corrigido anteriormente
        f_transp = c_file.file_uploader("PDF das Transparências", type="pdf")
        if c_act.button("✨ Gerar Texto (IA)", type="primary", use_container_width=True):
            if not f_transp: st.warning("Adicione um PDF.")
            else:
                with st.spinner("Processando..."):
                    txt = ia_utils.extract_text_from_pdf(f_transp)
                    res = ia_utils.process_transparencies(txt, "")
                    if "Erro" not in res:
                        st.session_state.data_store["transparencias_data"] = res
                        st.rerun()
                    else: st.error(res["Erro"])
        
        diretorias_ignorar = st.multiselect("Diretorias sem apresentação (ignorar):", options=config.DIRETORIAS)
        if diretorias_ignorar and st.session_state.data_store.get("transparencias_data"):
            for d in diretorias_ignorar:
                if d in st.session_state.data_store["transparencias_data"]:
                    st.session_state.data_store["transparencias_data"][d] = {"Realizado": "Não foram apresentadas atividades.", "Planejado": "Não foram apresentadas atividades."}
        
        # Editor Visual
        data_t = st.session_state.data_store.get("transparencias_data", {})
        if data_t and "Erro" not in data_t:
            st.subheader("📝 Revisão")
            for diretoria in config.DIRETORIAS:
                if diretoria not in data_t: continue
                with st.expander(f"🔹 {diretoria}"):
                    data_t[diretoria]["Realizado"] = st.text_area(f"R - {diretoria}", value=data_t[diretoria].get("Realizado", ""), height=100)
                    data_t[diretoria]["Planejado"] = st.text_area(f"P - {diretoria}", value=data_t[diretoria].get("Planejado", ""), height=100)

# --- TAB 3: PAUTAS ---
with tab3:
    st.header("Pautas")
    
    # Inicializa lista se vazia
    if not st.session_state.pautas_dinamicas:
        st.info("Nenhuma pauta iniciada. Clique em '➕ Nova Pauta' abaixo.")

    pautas_rem = []
    
    for i, pauta in enumerate(st.session_state.pautas_dinamicas):
        # Garante ID único para widgets
        pauta_id = pauta['id']
        key_textarea = f"tf_{pauta_id}"
        
        titulo_display = pauta['titulo'] if pauta['titulo'] else f"Pauta {i+1}"
        
        with st.expander(f"📂 {titulo_display}", expanded=True):
            # Layout Título e Lixeira
            c1, c2 = st.columns([5, 0.5], vertical_alignment="bottom")
            new_tit = c1.text_input("Título", value=pauta['titulo'], key=f"t_{pauta_id}")
            st.session_state.pautas_dinamicas[i]['titulo'] = new_tit
            
            if c2.button("🗑️", key=f"d_{pauta_id}"): 
                pautas_rem.append(i)
            
            # Layout PDF e Notas
            c_up, c_nt = st.columns(2)
            pdf = c_up.file_uploader("PDF de Apoio", type="pdf", key=f"p_{pauta_id}")
            new_note = c_nt.text_area("Notas de Apoio", value=pauta['notas'], key=f"n_{pauta_id}", height=100)
            st.session_state.pautas_dinamicas[i]['notas'] = new_note
            
            # --- CORREÇÃO DO BOTÃO DE GERAÇÃO ---
            if st.button(f"⚡ Gerar Texto (Pauta {i+1})", key=f"g_{pauta_id}", type="primary"):
                with st.spinner("Lendo documento e escrevendo ata..."):
                    # 1. Extração
                    txt_extracted = ia_utils.extract_text_from_pdf(pdf) if pdf else ""
                    
                    # 2. Geração IA
                    res = ia_utils.process_pauta(
                        txt_extracted, 
                        pauta['notas'], 
                        pauta['titulo'], 
                        history_txt
                    )
                    
                    # 3. ATUALIZAÇÃO FORÇADA (O Pulo do Gato)
                    # Atualiza o dado na lista (para o documento final)
                    st.session_state.pautas_dinamicas[i]['texto'] = res
                    # Atualiza o widget visual (para a aba 3)
                    st.session_state[key_textarea] = res
                    
                    st.success("Gerado!")
                    st.rerun()
            
            # Área de Texto Final (Bidirecional)
            # Lê o valor atual da lista, mas permite edição manual
            val_atual = st.session_state.pautas_dinamicas[i]['texto']
            
            texto_final_usuario = st.text_area(
                "Texto Final (Editável)", 
                value=val_atual, 
                key=key_textarea,
                height=200
            )
            
            # Salva qualquer edição manual do usuário de volta na lista
            st.session_state.pautas_dinamicas[i]['texto'] = texto_final_usuario

    # Processa exclusões fora do loop visual
    if pautas_rem:
        for i in sorted(pautas_rem, reverse=True): 
            del st.session_state.pautas_dinamicas[i]
        st.rerun()
    
    st.markdown("---")
    if st.button("➕ Nova Pauta"):
        st.session_state.pautas_dinamicas.append({
            "id": str(uuid.uuid4()), 
            "titulo": "", 
            "texto": "", 
            "notas": ""
        })
        st.rerun()

# --- TAB 4: SEÇÕES FINAIS ---
with tab4:
    st.header("Seções Finais")
    
    # Avisos
    st.subheader("⚠️ Avisos Gerais")
    chk_avisos = st.checkbox("Houve Avisos?", value=False)
    if chk_avisos:
        c1, c2 = st.columns([4, 1], vertical_alignment="center")
        av_txt = c1.text_area("Texto dos Avisos", value=st.session_state.data_store.get("avisos_text", ""))
        if c2.button("Padronizar (IA)"):
            av_txt = ia_utils.refine_notices(av_txt)
            st.session_state.data_store["avisos_text"] = av_txt
            st.rerun()
        st.session_state.data_store["avisos_text"] = av_txt
    else:
        st.session_state.data_store["avisos_text"] = ""

    st.markdown("---")
    
    # Sparckselt
    st.subheader("💎 Sparckselt")
    chk_sparck = st.checkbox("Houve passagem do Sparckselt?", value=False)
    if chk_sparck:
        c1, c2 = st.columns(2)
        ant = c1.text_input("Quem passou?")
        nov = c2.text_input("Quem recebeu?")
        mot = st.text_area("Motivo")
        st.session_state.data_store["sparckselt"] = {"anterior": ant, "novo": nov, "motivo": mot}
    else:
        st.session_state.data_store["sparckselt"] = {}

    st.markdown("---")

    # Elogios
    st.subheader("💌 Caixinha de Elogios")
    chk_elogios = st.checkbox("Houve leitura de Elogios?", value=False)
    if chk_elogios:
        nome = st.text_input("Responsável pela leitura")
        st.session_state.data_store["elogios_leitor"] = nome
    else:
        st.session_state.data_store["elogios_leitor"] = ""

# --- TAB 5: REVISÃO E CONCLUSÃO ---
with tab5:
    st.header("Revisão e Conclusão")
    
    # 1. Auditoria Rápida
    st.markdown("### 🤖 Auditoria Rápida")
    if st.button("Analisar Inconsistências"):
        context = {
            "pautas": st.session_state.pautas_dinamicas,
            "transparencias": st.session_state.data_store.get("transparencias_data", {}),
            "avisos": st.session_state.data_store.get("avisos_text", "")
        }
        with st.spinner("Auditando..."):
            audit = ia_utils.audit_meeting_summary(context)
            st.info(audit)
            
            # Botão de Aplicar Correções (Só aparece após auditar)
            st.session_state['show_apply_fix'] = True

    # Botão de Aplicação de Correções
    if st.session_state.get('show_apply_fix'):
        st.markdown("#### 🛠️ Correção Automática")
        st.caption("A IA irá reescrever os textos armazenados para corrigir gramática e tom.")
        if st.button("Aplicar Correções Sugeridas pela IA"):
            with st.spinner("Aplicando melhorias nos textos..."):
                context = {
                    "pautas": st.session_state.pautas_dinamicas,
                    "transparencias": st.session_state.data_store.get("transparencias_data", {}),
                    "avisos": st.session_state.data_store.get("avisos_text", "")
                }
                new_data = ia_utils.apply_auto_corrections(context)
                
                if "Erro" not in new_data:
                    if "pautas" in new_data:
                        st.session_state.pautas_dinamicas = merge_corrected_pautas(
                            st.session_state.pautas_dinamicas,
                            new_data["pautas"],
                        )
                    if "transparencias" in new_data:
                        st.session_state.data_store["transparencias_data"] = new_data["transparencias"]
                    if "avisos" in new_data:
                        st.session_state.data_store["avisos_text"] = new_data["avisos"]
                    st.success("Textos atualizados com sucesso! Verifique nas abas anteriores.")
                    st.rerun()
                else:
                    st.error(f"Falha ao aplicar correções: {new_data['Erro']}")

    st.markdown("---")

    # 2. Geração e PDF
    st.markdown("### 📄 Documento Final")
    col_gen, col_down = st.columns(2)
    
    filename_base = generate_filename() # Ex: ATA RG - 12/05.docx
    active_template_path = config.get_active_ata_template_path()
    st.caption(f"Template ativo: {active_template_path}")
    
    if col_gen.button("🚀 Gerar Documento", type="primary"):
        # Contexto Jinja
        ctx = {
            **st.session_state.data_store["meta"],
            "lista_pautas": st.session_state.pautas_dinamicas,
            "texto_avisos": st.session_state.data_store.get("avisos_text", ""),
            "texto_sparckselt": bool(st.session_state.data_store.get("sparckselt", {}).get("novo")),
            "s": st.session_state.data_store.get("sparckselt", {}),
            "texto_elogios": bool(st.session_state.data_store.get("elogios_leitor")),
            "ce": {"nomes": st.session_state.data_store.get("elogios_leitor")},
            "tem_transparencias": bool(st.session_state.data_store.get("transparencias_data"))
        }
        
        # Flattening Transparências
        raw_t = st.session_state.data_store.get("transparencias_data", {})
        raw_t_norm = {k.lower(): v for k, v in raw_t.items()}
        for k, v in config.MAP_JINJA.items():
            if k in raw_t_norm:
                ctx[f"tr_{v}"] = True
                ctx[f"tr_{v}_realizadas"] = raw_t_norm[k]["Realizado"]
                ctx[f"tr_{v}_planejadas"] = raw_t_norm[k]["Planejado"]

        try:
            doc = DocxTemplate(active_template_path)
            doc.render(ctx)
            bio = io.BytesIO()
            doc.save(bio)

            docx_bytes = bio.getvalue()
            set_final_document_state(
                file_bytes=docx_bytes,
                filename=filename_base,
                file_type="docx",
                docx_bytes=docx_bytes,
                pdf_bytes=None,
            )
            st.success(f"Gerado: {filename_base}")
            
            # Tentativa de converter para PDF (Só funciona se tiver Office/LibreOffice no servidor)
            if HAS_PDF_CONVERTER:
                try:
                    with open(config.TEMP_DOCX_PATH, "wb") as f: f.write(docx_bytes)
                    convert(config.TEMP_DOCX_PATH, config.TEMP_PDF_PATH)
                    with open(config.TEMP_PDF_PATH, "rb") as f:
                        pdf_bytes = f.read()
                        set_final_document_state(
                            file_bytes=pdf_bytes,
                            filename=filename_base.replace(".docx", ".pdf"),
                            file_type="pdf",
                            docx_bytes=docx_bytes,
                            pdf_bytes=pdf_bytes,
                        )
                    st.info("Convertido para PDF automaticamente.")
                except:
                    st.warning("Não foi possível converter para PDF automaticamente (ambiente sem Office). Mantendo .docx.")
            else:
                 st.caption("Ambiente Linux/Cloud: Conversão nativa PDF desativada. O arquivo será .DOCX (Authentique aceita).")

        except Exception as e:
            st.error(f"Erro: {e}")

    # Download
    if st.session_state.get('final_file_blob'):
        col_down.download_button(
            f"📥 Baixar {st.session_state['final_filename']}", 
            st.session_state['final_file_blob'], 
            st.session_state['final_filename']
        )

    preview_pdf = st.session_state.get('final_pdf_blob')
    preview_docx = st.session_state.get('final_docx_blob')
    if preview_pdf or preview_docx:
        st.markdown("#### Visualização do documento")
        preview_mode_options = ["Documento .docx"]
        if preview_pdf:
            preview_mode_options.insert(0, "PDF")

        preview_mode = st.radio(
            "Formato de visualização",
            options=preview_mode_options,
            horizontal=True,
            key="ata_preview_mode",
        )

        if preview_mode == "PDF" and preview_pdf:
            render_pdf_preview(preview_pdf)
        elif preview_docx:
            render_docx_preview(preview_docx)
            if not preview_pdf:
                st.info("A hospedagem atual não converte DOCX para PDF. A revisão está usando a visualização direta do documento .docx.")
    
    st.markdown("---")
    st.markdown("**Validação Humana:**")
    opt = st.radio("Como prosseguir?", ["Usar arquivo gerado acima", "Fazer upload de versão corrigida (PDF)"])
    
    if opt == "Fazer upload de versão corrigida (PDF)":
        up = st.file_uploader("Upload PDF Final", type="pdf")
        if up:
            pdf_bytes = up.getvalue()
            set_final_document_state(
                file_bytes=pdf_bytes,
                filename=up.name,
                file_type="pdf",
                docx_bytes=st.session_state.get('final_docx_blob'),
                pdf_bytes=pdf_bytes,
            )
            st.success("Arquivo manual carregado.")

# --- TAB 6: ASSINATURA ---
with tab6:
    st.header("Assinatura Oficial")
    
    # Verifica arquivo
    file_blob = st.session_state.get('final_file_blob')
    if not file_blob:
        st.warning("⚠️ Gere ou suba o documento na aba anterior (Revisão e Conclusão).")
        st.stop()
        
    # Verifica lista
    raw_att = st.session_state.data_store.get("attendance_raw", "")
    if not raw_att:
        st.error("⚠️ Lista de presença vazia! Volte na Aba 1.")
        st.stop()

    # Info do Documento
    deadline_str = st.session_state.get("final_deadline", "Calculando...")
    filename = st.session_state.get('final_filename', 'ATA.docx')
    
    st.info(f"📂 Arquivo: **{filename}**\n\n📅 Prazo: **{deadline_str}**")

    # 1. Conferência
    st.markdown("### 1. Conferência de Signatários")
    if st.button("🔍 Analisar E-mails"):
        signers, missing, display_map = authentique_utils.get_signers_emails(raw_att)
        st.session_state['auth_preview'] = {
            "signers": signers, 
            "missing": missing, 
            "map": display_map
        }

    # 2. Envio
    if 'auth_preview' in st.session_state:
        prev = st.session_state['auth_preview']
        
        df = pd.DataFrame(prev['map'], columns=["Nome", "Email", "Status"])
        st.table(df)
        
        if prev['missing']:
            st.warning(f"⚠️ {len(prev['missing'])} pessoas sem e-mail não receberão o convite.")
        
        st.markdown("### 2. Disparo")
        if st.button("✅ Confirmar e Enviar para Authentique", type="primary"):
            file_mem = io.BytesIO(file_blob)
            file_mem.name = filename
            
            with st.spinner("Conectando à API Authentique..."):
                try:
                    doc_id = authentique_utils.send_to_authentique(file_mem, prev['signers'], doc_name=filename)
                    st.success(f"🎉 Documento enviado!")    
                    # Limpa preview para evitar reenvio
                    del st.session_state['auth_preview']
                except Exception as e:
                    st.error(f"{e}")

# --- TAB 7: NOTIFICAÇÃO ---
with tab7:
    st.header("Notificação por Email")
    
    prazo_txt = st.session_state.get("final_deadline", "DD/MM/AAAA")
    
    default_body = f"""Bom dia, boa tarde e boa noite CONSELT!
[Mensagem personalizada]. Qualquer problema, chama JF! 🔥🔥

[AQUI ENTRA A GIF]

Foi enviado um e-mail para todos os efetivos presentes para assinar a ATA.
A assinatura é obrigatória e tem como prazo até o dia {prazo_txt}."""

    with st.form("email_form"):
        # Correção de Layout: Input full width fora de colunas apertadas
        email_to = st.text_input("Para (separe por vírgula):", value="conselt@conselt.com.br, trainees@conselt.com.br")
        subject = st.text_input("Assunto:", value="Assinatura da ATA - Reunião Geral")
        body_content = st.text_area("Corpo do Email", value=default_body, height=200)
        
        # Correção: O Uploader ficava feio sem label
        st.markdown("**Anexar GIF/Imagem:**")
        gif_file = st.file_uploader("Mídia Visual", type=["png", "jpg", "gif"])
        
        if st.form_submit_button("📧 Disparar Emails"):
            final_file_blob = st.session_state.get('final_file_blob')
            if not final_file_blob:
                st.error("Nenhum documento de ATA disponível.")
            else:
                try:
                    # Passando os valores CORRETOS dos inputs, não hardcoded
                    recipients = [e.strip() for e in email_to.split(",") if e.strip()]
                    
                    file_mem = io.BytesIO(final_file_blob)
                    
                    email_utils.send_notification_email(
                        file_obj=file_mem,
                        filename=st.session_state.get('final_filename', 'ATA.docx'),
                        receivers_list=recipients, # Usa a lista do input
                        subject=subject,
                        body_text=body_content
                    )
                    st.success(f"Email enviado para: {recipients}")
                except Exception as e:
                    st.error(f"Erro: {e}")